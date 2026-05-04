import json
import os
import platform
import re
import shutil
import subprocess
from datetime import datetime, timedelta
from pathlib import Path

import requests
import tkinter as tk
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from tkcalendar import DateEntry
from tkinter import filedialog, messagebox
from tkinter import ttk


APP_NAME = "Schuljahresplan Agent"
APP_VERSION = "1.1"
SETTINGS_FILE = "settings.json"

BASE_PATH = Path(os.getcwd())
SETTINGS_PATH = BASE_PATH / SETTINGS_FILE

AUSGABE_ORDNER = BASE_PATH / "ausgabe"
BACKUP_ORDNER = AUSGABE_ORDNER / "backups"
STANDARD_AUSGABE = AUSGABE_ORDNER / "Jahresarbeitsplan_neu.docx"
aktueller_plan = STANDARD_AUSGABE

BUNDESLAENDER = {
    "Sachsen": ("SN", "DE-SN"),
    "Bayern": ("BY", "DE-BY"),
    "Berlin": ("BE", "DE-BE"),
    "Brandenburg": ("BB", "DE-BB"),
    "Baden-Württemberg": ("BW", "DE-BW"),
    "Bremen": ("HB", "DE-HB"),
    "Hamburg": ("HH", "DE-HH"),
    "Hessen": ("HE", "DE-HE"),
    "Mecklenburg-Vorpommern": ("MV", "DE-MV"),
    "Niedersachsen": ("NI", "DE-NI"),
    "Nordrhein-Westfalen": ("NW", "DE-NW"),
    "Rheinland-Pfalz": ("RP", "DE-RP"),
    "Saarland": ("SL", "DE-SL"),
    "Sachsen-Anhalt": ("ST", "DE-ST"),
    "Schleswig-Holstein": ("SH", "DE-SH"),
    "Thüringen": ("TH", "DE-TH"),
}

SCHULJAHRE = [
    "2025/2026",
    "2026/2027",
    "2027/2028",
    "2028/2029",
    "2029/2030",
    "2030/2031",
    "2031/2032",
]

SPALTEN = {0: 2, 1: 3, 2: 4, 3: 5, 4: 6}

FARBE_HEADER = "D9EAF7"
FARBE_FERIEN = "BFBFBF"
FARBE_FREI = "FFF2CC"
FARBE_WEISS = "FFFFFF"


# =========================
# Einstellungen
# =========================

def load_settings():
    if not SETTINGS_PATH.exists():
        return {}

    try:
        with open(SETTINGS_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def save_settings(data):
    try:
        with open(SETTINGS_PATH, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=4, ensure_ascii=False)
    except Exception as e:
        messagebox.showwarning("Hinweis", f"Einstellungen konnten nicht gespeichert werden:\n{e}")


def save_current_settings():
    data = {
        "bundesland": bundesland_var.get(),
        "schuljahr": schuljahr_var.get(),
        "frei_bewegliche_ferientage": frei_input.get().strip(),
        "letzter_plan": str(aktueller_plan) if aktueller_plan else "",
    }
    save_settings(data)


# =========================
# Word-Hilfsfunktionen
# =========================

def backup_datei(pfad):
    if not pfad.exists():
        return

    BACKUP_ORDNER.mkdir(parents=True, exist_ok=True)

    zeit = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    backup_name = f"{pfad.stem}_{zeit}{pfad.suffix}"
    backup_pfad = BACKUP_ORDNER / backup_name

    shutil.copy2(pfad, backup_pfad)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    for shd in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(shd)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_bold(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True


def set_cell_text(cell, text, bold=False, font_size=8, align="center"):
    cell.text = ""
    p = cell.paragraphs[0]

    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(8)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_column_width(table, column_index, width_cm):
    for row in table.rows:
        cell = row.cells[column_index]
        cell.width = Cm(width_cm)
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        tc_w = tc_pr.first_child_found_in("w:tcW")
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            tc_pr.append(tc_w)
        tc_w.set(qn("w:w"), str(int(width_cm * 567)))
        tc_w.set(qn("w:type"), "dxa")


def clear_cell(cell):
    set_cell_text(cell, "", font_size=8, align="left")
    set_cell_shading(cell, FARBE_WEISS)


def add_text(cell, text):
    if text in cell.text:
        return

    if cell.text.strip():
        cell.text = cell.text.strip() + "\n" + text
    else:
        cell.text = text

    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.size = Pt(7)


def erstelle_leeres_dokument(schuljahr, bundesland):
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"Schuljahresplan {schuljahr}")
    title_run.bold = True
    title_run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"Bundesland: {bundesland}")
    subtitle_run.font.size = Pt(10)

    doc.add_paragraph("")

    # 1 Kopfzeile + 45 Wochenzeilen reichen für ein Schuljahr sicher aus.
    table = doc.add_table(rows=46, cols=7)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = [
    "KW",
    "Zeitraum",
    "Montag",
    "Dienstag",
    "Mittwoch",
    "Donnerstag",
    "Freitag"
   ]
    header_cells = table.rows[0].cells

    for index, header in enumerate(headers):
        cell = header_cells[index]
        set_cell_text(cell, header, bold=True, font_size=9, align="center")
        set_cell_shading(cell, FARBE_HEADER)
        set_cell_bold(cell)

        widths = [1.2, 4.5, 4.5, 4.5, 4.5, 4.5, 4.5]
    for index, width in enumerate(widths):
        set_column_width(table, index, width)

    return doc, table


# =========================
# Datums-/API-Funktionen
# =========================

def parse_schuljahr(text):
    match = re.match(r"^(\d{4})/(\d{4})$", text.strip())
    if not match:
        return None, None

    start = int(match.group(1))
    ende = int(match.group(2))

    if ende != start + 1:
        return None, None

    return start, ende


def parse_date(text):
    text = text.strip()[:10]
    if "-" in text:
        return datetime.strptime(text, "%Y-%m-%d").date()
    return datetime.strptime(text, "%d.%m.%Y").date()


def format_de(d):
    return d.strftime("%d.%m.%Y")


def extract_name(item):
    names = item.get("name") or item.get("localName") or []

    if isinstance(names, list) and names:
        return names[0].get("text", "Ferien")

    if isinstance(names, str):
        return names

    return "Ferien"


def lade_schulferien(startjahr, endjahr, subdivision_code):
    url = "https://openholidaysapi.org/SchoolHolidays"
    params = {
        "countryIsoCode": "DE",
        "subdivisionCode": subdivision_code,
        "languageIsoCode": "DE",
        "validFrom": f"{startjahr}-08-01",
        "validTo": f"{endjahr}-08-31",
    }

    r = requests.get(url, params=params, timeout=15)
    r.raise_for_status()

    eintraege = []

    for item in r.json():
        name = extract_name(item)
        start = parse_date(item["startDate"])
        ende = parse_date(item["endDate"])

        typ = "ferien"
        if "unterrichtsfreier tag" in name.lower():
            typ = "frei"

        eintraege.append({
            "name": name,
            "start": start,
            "ende": ende,
            "typ": typ,
        })

    return eintraege


def lade_feiertage(startjahr, endjahr, land_code):
    alle = []

    for jahr in [startjahr, endjahr]:
        url = f"https://feiertage-api.de/api/?jahr={jahr}&nur_land={land_code}"
        r = requests.get(url, timeout=15)
        r.raise_for_status()

        for name, info in r.json().items():
            d = parse_date(info["datum"])

            if name == "Fronleichnam" and land_code == "SN":
                continue

            alle.append({
                "name": name,
                "start": d,
                "ende": d,
                "typ": "feiertag",
            })

    return alle


def validiere_frei_tage(text, eintraege):
    tage = []

    if not text.strip():
        return tage

    for teil in text.split(","):
        datum_str = teil.strip()

        if not datum_str:
            continue

        try:
            d = parse_date(datum_str)
        except ValueError:
            messagebox.showerror("Fehler", f"Ungültiges Datum:\n{datum_str}")
            return None

        if d.weekday() > 4:
            messagebox.showerror("Fehler", f"{datum_str} liegt am Wochenende.")
            return None

        for e in eintraege:
            if e["start"] <= d <= e["ende"]:
                messagebox.showerror(
                    "Fehler",
                    f"{datum_str} liegt bereits auf:\n{e['name']}\n\nBitte anderen Tag wählen."
                )
                return None

        tage.append({
            "name": "Frei beweglicher Ferientag",
            "start": d,
            "ende": d,
            "typ": "frei",
        })

    return tage


def erster_montag_nach_sommerferien(eintraege, startjahr):
    sommerferien = [
        e for e in eintraege
        if e["typ"] == "ferien"
        and "sommer" in e["name"].lower()
        and e["start"].year == startjahr
    ]

    if not sommerferien:
        return None

    ende = max(sommerferien, key=lambda x: x["ende"])["ende"]
    tag = ende + timedelta(days=1)

    while tag.weekday() != 0:
        tag += timedelta(days=1)

    return tag


def eintraege_fuer_tag(eintraege, tag):
    passende = [e for e in eintraege if e["start"] <= tag <= e["ende"]]

    ferien = [e for e in passende if e["typ"] == "ferien"]
    frei = [e for e in passende if e["typ"] == "frei"]
    feiertage = [e for e in passende if e["typ"] == "feiertag"]

    if ferien:
        return ferien
    if frei:
        return frei
    return feiertage


# =========================
# Hauptfunktionen
# =========================

def create_plan():
    global aktueller_plan

    schuljahr = schuljahr_var.get()
    startjahr, endjahr = parse_schuljahr(schuljahr)

    if not startjahr:
        messagebox.showerror("Fehler", "Ungültiges Schuljahr.")
        return

    bundesland = bundesland_var.get()
    land_code, subdivision_code = BUNDESLAENDER[bundesland]

    try:
        status_var.set("Lade Ferien und Feiertage ...")
        root.update_idletasks()

        eintraege = []
        eintraege += lade_schulferien(startjahr, endjahr, subdivision_code)
        eintraege += lade_feiertage(startjahr, endjahr, land_code)

        frei = validiere_frei_tage(frei_input.get(), eintraege)
        if frei is None:
            status_var.set("Bereit.")
            return

        eintraege += frei

        start_montag = erster_montag_nach_sommerferien(eintraege, startjahr)

        if start_montag is None:
            messagebox.showerror("Fehler", "Erster Schultag konnte nicht berechnet werden.")
            status_var.set("Fehler.")
            return

        doc, table = erstelle_leeres_dokument(schuljahr, bundesland)

        woche = start_montag

        for row in table.rows[1:]:
            montag = woche
            freitag = montag + timedelta(days=4)

            set_cell_text(row.cells[0], str(montag.isocalendar().week), font_size=7, align="center")
            set_cell_text(row.cells[1], f"{format_de(montag)} – {format_de(freitag)}", font_size=7, align="center")

            for i in range(2, 7):
                clear_cell(row.cells[i])

            for offset in range(5):
                tag = montag + timedelta(days=offset)
                cell = row.cells[2 + offset]

                for e in eintraege_fuer_tag(eintraege, tag):
                    add_text(cell, e["name"])

                    if e["typ"] == "ferien":
                        set_cell_shading(cell, FARBE_FERIEN)
                    elif e["typ"] == "frei":
                        set_cell_shading(cell, FARBE_FREI)
                    elif e["typ"] == "feiertag":
                        set_cell_bold(cell)

            woche += timedelta(days=7)

        safe_schuljahr = schuljahr.replace("/", "-")
        dateiname = f"Schuljahresplan {safe_schuljahr}.docx"
        ziel = AUSGABE_ORDNER / dateiname

        ziel.parent.mkdir(exist_ok=True)
        backup_datei(ziel)
        doc.save(ziel)

        aktueller_plan = ziel
        plan_label.config(text=f"Aktueller Plan:\n{aktueller_plan}")

        save_current_settings()

        status_var.set(f"Plan erstellt: {dateiname}")
        messagebox.showinfo("OK", f"Plan erstellt:\n{ziel}")

    except requests.RequestException as e:
        status_var.set("Fehler beim Abrufen der Daten.")
        messagebox.showerror("Fehler", f"Ferien oder Feiertage konnten nicht abgerufen werden:\n{e}")
    except Exception as e:
        status_var.set("Fehler beim Erstellen.")
        messagebox.showerror("Fehler", f"Plan konnte nicht erstellt werden:\n{e}")


def lade_plan():
    global aktueller_plan

    datei = filedialog.askopenfilename(
        title="Vorhandenen Schuljahresplan auswählen",
        filetypes=[("Word-Dateien", "*.docx")]
    )

    if not datei:
        return

    aktueller_plan = Path(datei)
    plan_label.config(text=f"Aktueller Plan:\n{aktueller_plan}")
    status_var.set("Plan geladen.")

    save_current_settings()

    messagebox.showinfo("OK", "Plan geladen.")


def parse_event_zeit(text):
    text = text.strip()

    if not text:
        return None

    if not re.match(r"^\d{1,2}[:.]\d{2}$", text):
        messagebox.showerror("Fehler", "Uhrzeit bitte im Format 08:00 oder 8.00 eingeben.")
        return False

    text = text.replace(".", ":")
    stunde, minute = text.split(":")

    stunde = int(stunde)
    minute = int(minute)

    if stunde > 23 or minute > 59:
        messagebox.showerror("Fehler", "Bitte eine gültige Uhrzeit eingeben.")
        return False

    return f"{stunde:02d}:{minute:02d}"


def sortiere_events_in_zelle(cell):
    zeilen = [z.strip() for z in cell.text.split("\n") if z.strip()]

    def sort_key(zeile):
        match = re.match(r"^(\d{2}):(\d{2})", zeile)
        if match:
            return (0, int(match.group(1)), int(match.group(2)), zeile)
        return (1, 99, 99, zeile)

    cell.text = "\n".join(sorted(set(zeilen), key=sort_key))

    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.size = Pt(7)


def event_hinzufuegen():
    if not aktueller_plan.exists():
        messagebox.showerror("Fehler", "Bitte zuerst einen Plan erstellen oder laden.")
        return

    datum_event = event_datum.get_date()
    ereignis = event_text.get().strip()
    zeit = event_zeit.get().strip()

    if not ereignis:
        messagebox.showerror("Fehler", "Bitte Ereignis eintragen.")
        return

    zeit_formatiert = parse_event_zeit(zeit)
    if zeit_formatiert is False:
        return

    eintrag = f"{zeit_formatiert} {ereignis}" if zeit_formatiert else ereignis

    if datum_event.weekday() > 4:
        messagebox.showerror("Fehler", "Das Datum liegt am Wochenende.")
        return

    try:
        backup_datei(aktueller_plan)

        doc = Document(aktueller_plan)
        table = doc.tables[0]

        ziel_spalte = SPALTEN[datum_event.weekday()]

        for row in table.rows[1:]:
            zeitraum = row.cells[3].text.strip()

            if "–" not in zeitraum:
                continue

            start_text, ende_text = [x.strip() for x in zeitraum.split("–")]

            try:
                start = parse_date(start_text)
                ende = parse_date(ende_text)
            except ValueError:
                continue

            if start <= datum_event <= ende:
                zelle = row.cells[ziel_spalte]

                if eintrag in zelle.text:
                    messagebox.showwarning("Doppelt", "Dieses Ereignis steht dort bereits.")
                    return

                add_text(zelle, eintrag)
                sortiere_events_in_zelle(zelle)

                doc.save(aktueller_plan)

                event_text.delete(0, tk.END)
                event_zeit.delete(0, tk.END)

                save_current_settings()

                status_var.set("Ereignis hinzugefügt. Backup erstellt.")
                messagebox.showinfo("OK", "Ereignis hinzugefügt.\nBackup wurde erstellt.")
                return

        messagebox.showerror("Fehler", "Passende Woche nicht gefunden.")

    except Exception as e:
        status_var.set("Fehler beim Hinzufügen.")
        messagebox.showerror("Fehler", f"Ereignis konnte nicht hinzugefügt werden:\n{e}")


def plan_oeffnen():
    if not aktueller_plan.exists():
        messagebox.showerror("Fehler", "Kein Plan vorhanden.")
        return

    system = platform.system()

    try:
        if system == "Windows":
            os.startfile(str(aktueller_plan))
        elif system == "Linux":
            subprocess.Popen(["xdg-open", str(aktueller_plan)])
        elif system == "Darwin":
            subprocess.Popen(["open", str(aktueller_plan)])
        else:
            messagebox.showerror("Fehler", f"Nicht unterstütztes System: {system}")
    except Exception as e:
        messagebox.showerror("Fehler", f"Plan konnte nicht geöffnet werden:\n{e}")


def update_datum_label(event=None):
    d = event_datum.get_date()
    tage = ["Mo", "Di", "Mi", "Do", "Fr", "Sa", "So"]
    auswahl_label.config(text=f"{tage[d.weekday()]} {d.strftime('%d.%m.%Y')}")


# =========================
# GUI
# =========================

settings = load_settings()

saved_bundesland = settings.get("bundesland", "Sachsen")
if saved_bundesland not in BUNDESLAENDER:
    saved_bundesland = "Sachsen"

saved_schuljahr = settings.get("schuljahr", "2026/2027")
if saved_schuljahr not in SCHULJAHRE:
    saved_schuljahr = "2026/2027"

saved_frei_tage = settings.get("frei_bewegliche_ferientage", "")

saved_plan = settings.get("letzter_plan", "")
if saved_plan and Path(saved_plan).exists():
    aktueller_plan = Path(saved_plan)

root = tk.Tk()
root.title(f"{APP_NAME} {APP_VERSION}")
root.minsize(900, 800)

# Linux-kompatible Maximierung. Falls das System es nicht unterstützt, wird eine große Fenstergröße verwendet.
try:
    root.attributes("-zoomed", True)
except tk.TclError:
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()
    root.geometry(f"{screen_width}x{screen_height}+0+0")

style = ttk.Style()
style.theme_use("clam")

style.configure("TFrame", background="#f4f6fb")
style.configure("Card.TFrame", background="#ffffff", relief="flat")
style.configure("TLabel", background="#f4f6fb", foreground="#1f2937", font=("Arial", 10))
style.configure("Title.TLabel", background="#f4f6fb", foreground="#111827", font=("Arial", 18, "bold"))
style.configure("SubTitle.TLabel", background="#ffffff", foreground="#111827", font=("Arial", 13, "bold"))
style.configure("Hint.TLabel", background="#ffffff", foreground="#6b7280", font=("Arial", 9))
style.configure("CardText.TLabel", background="#ffffff", foreground="#1f2937", font=("Arial", 10))
style.configure("Status.TLabel", background="#e5e7eb", foreground="#374151", font=("Arial", 9))
style.configure("TButton", font=("Arial", 10, "bold"), padding=8)
style.configure("Primary.TButton", font=("Arial", 10, "bold"), padding=9)
style.configure("TEntry", padding=5)
style.configure("TCombobox", padding=5)

main = ttk.Frame(root, padding=18)
main.pack(fill="both", expand=True)

ttk.Label(main, text=APP_NAME, style="Title.TLabel").pack(anchor="center", pady=(0, 4))
ttk.Label(
    main,
    text=f"Version {APP_VERSION} · erstellt Schuljahrespläne ohne externe Word-Vorlage · © 2026 Max K.",
    style="TLabel"
).pack(anchor="center", pady=(0, 14))

card_create = ttk.Frame(main, style="Card.TFrame", padding=18)
card_create.pack(fill="x", pady=10)

ttk.Label(card_create, text="Neuen Plan erstellen", style="SubTitle.TLabel").grid(row=0, column=0, columnspan=2, sticky="w", pady=(0, 12))

ttk.Label(card_create, text="Bundesland", style="CardText.TLabel").grid(row=1, column=0, sticky="w", pady=5)
bundesland_var = tk.StringVar(value=saved_bundesland)
bundesland_box = ttk.Combobox(card_create, textvariable=bundesland_var, values=list(BUNDESLAENDER.keys()), state="readonly", width=28)
bundesland_box.grid(row=1, column=1, sticky="ew", pady=5)

ttk.Label(card_create, text="Schuljahr", style="CardText.TLabel").grid(row=2, column=0, sticky="w", pady=5)
schuljahr_var = tk.StringVar(value=saved_schuljahr)
schuljahr_box = ttk.Combobox(card_create, textvariable=schuljahr_var, values=SCHULJAHRE, state="readonly", width=28)
schuljahr_box.grid(row=2, column=1, sticky="ew", pady=5)

ttk.Label(card_create, text="Frei bewegliche Ferientage", style="CardText.TLabel").grid(row=3, column=0, sticky="w", pady=5)
frei_input = ttk.Entry(card_create, width=45)
frei_input.grid(row=3, column=1, sticky="ew", pady=5)
frei_input.insert(0, saved_frei_tage)

ttk.Label(card_create, text="Mehrere Daten mit Komma trennen, z.B. 24.05.2027", style="Hint.TLabel").grid(row=4, column=1, sticky="w", pady=(0, 10))

ttk.Button(card_create, text="Neuen Schuljahresplan erstellen", command=create_plan, style="Primary.TButton").grid(row=5, column=1, sticky="e", pady=(8, 0))

card_create.columnconfigure(1, weight=1)

card_edit = ttk.Frame(main, style="Card.TFrame", padding=18)
card_edit.pack(fill="x", pady=10)

ttk.Label(card_edit, text="Aktueller Schuljahresplan", style="SubTitle.TLabel").pack(anchor="w", pady=(0, 10))

plan_label = ttk.Label(card_edit, text=f"Aktueller Plan:\n{aktueller_plan}", style="CardText.TLabel", wraplength=720)
plan_label.pack(anchor="w", pady=(0, 12))

button_row = ttk.Frame(card_edit, style="Card.TFrame")
button_row.pack(fill="x")

ttk.Button(button_row, text="Vorhandenen Plan laden", command=lade_plan).pack(side="left", padx=(0, 10))
ttk.Button(button_row, text="Aktuellen Plan öffnen", command=plan_oeffnen).pack(side="left")

card_event = ttk.Frame(main, style="Card.TFrame", padding=18)
card_event.pack(fill="x", pady=10)

ttk.Label(card_event, text="Ereignis in Plan eintragen", style="SubTitle.TLabel").grid(row=0, column=0, columnspan=2, sticky="w", pady=(0, 12))

ttk.Label(card_event, text="Datum", style="CardText.TLabel").grid(row=1, column=0, sticky="w", pady=5)
event_datum = DateEntry(card_event, date_pattern="dd.mm.yyyy", width=16)
event_datum.grid(row=1, column=1, sticky="w", pady=5)

auswahl_label = ttk.Label(card_event, text="", style="CardText.TLabel")
auswahl_label.grid(row=2, column=1, sticky="w", pady=(0, 8))
event_datum.bind("<<DateEntrySelected>>", update_datum_label)

ttk.Label(card_event, text="Uhrzeit optional", style="CardText.TLabel").grid(row=3, column=0, sticky="w", pady=5)
event_zeit = ttk.Entry(card_event, width=18)
event_zeit.grid(row=3, column=1, sticky="w", pady=5)

ttk.Label(card_event, text="Ereignis", style="CardText.TLabel").grid(row=4, column=0, sticky="w", pady=5)
event_text = ttk.Entry(card_event, width=55)
event_text.grid(row=4, column=1, sticky="ew", pady=5)

ttk.Button(card_event, text="Ereignis hinzufügen", command=event_hinzufuegen, style="Primary.TButton").grid(row=5, column=1, sticky="e", pady=(12, 0))

card_event.columnconfigure(1, weight=1)

status_var = tk.StringVar(value="Bereit · Einstellungen werden automatisch gespeichert")
status = ttk.Label(main, textvariable=status_var, style="Status.TLabel", padding=8)
status.pack(fill="x", pady=(12, 0))

update_datum_label()

root.mainloop()
